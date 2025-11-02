#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
增强版豆包提示词解析脚本
确保完整提取所有Word和文本文件中的提示词内容
"""

import os
import json
import re
from pathlib import Path
from docx import Document
import logging

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('parse_log.txt', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

class EnhancedPromptParser:
    def __init__(self, source_dir, output_file):
        self.source_dir = Path(source_dir)
        self.output_file = output_file
        self.prompts = []
        self.stats = {
            'total_files': 0,
            'docx_files': 0,
            'txt_files': 0,
            'parsed_successfully': 0,
            'failed_files': [],
            'categories': {}
        }
        
        # 分类映射表
        self.category_mapping = {
            '01自媒体类型': '自媒体',
            '02公文类': '公文写作',
            '03.英语类型': '英语学习',
            '04.论文类': '论文写作',
            '5.仿写类': '仿写创作',
            '6.影视小说类': '影视小说',
            '7.营销策划类': '营销策划',
            '8.职场类': '职场办公',
            '更新': '图像生成'  # 更新目录主要是图像生成相关
        }

    def scan_all_files(self):
        """扫描所有文件并返回文件列表"""
        files = []
        for ext in ['*.docx', '*.txt']:
            files.extend(self.source_dir.rglob(ext))
        
        self.stats['total_files'] = len(files)
        self.stats['docx_files'] = len([f for f in files if f.suffix == '.docx'])
        self.stats['txt_files'] = len([f for f in files if f.suffix == '.txt'])
        
        logging.info(f"扫描完成: 总文件数 {self.stats['total_files']}")
        logging.info(f"Word文档: {self.stats['docx_files']}, 文本文件: {self.stats['txt_files']}")
        
        return files

    def extract_from_docx(self, file_path):
        """从Word文档提取内容"""
        try:
            doc = Document(file_path)
            content_parts = []
            
            # 提取所有段落
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            content_parts.append(text)
            
            full_content = '\n'.join(content_parts)
            
            if not full_content.strip():
                logging.warning(f"Word文档内容为空: {file_path}")
                return None
                
            return full_content
            
        except Exception as e:
            logging.error(f"解析Word文档失败 {file_path}: {str(e)}")
            return None

    def extract_from_txt(self, file_path):
        """从文本文件提取内容"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read().strip()
                        if content:
                            return content
                except UnicodeDecodeError:
                    continue
            
            logging.error(f"无法解码文本文件: {file_path}")
            return None
            
        except Exception as e:
            logging.error(f"读取文本文件失败 {file_path}: {str(e)}")
            return None

    def determine_category(self, file_path):
        """根据文件路径确定分类"""
        path_parts = file_path.parts
        
        # 查找主分类目录
        for part in path_parts:
            for key, value in self.category_mapping.items():
                if key in part:
                    return value
        
        # 如果没有找到匹配的分类，根据文件名推断
        filename = file_path.name.lower()
        if any(word in filename for word in ['美女', '图像', '生成', 'cos', '绘本']):
            return '图像生成'
        elif any(word in filename for word in ['自媒体', '文案', '小红书', '公众号']):
            return '自媒体'
        elif any(word in filename for word in ['公文', 'ppt', '会议']):
            return '公文写作'
        elif any(word in filename for word in ['英语', '单词']):
            return '英语学习'
        elif any(word in filename for word in ['论文', '写作']):
            return '论文写作'
        elif any(word in filename for word in ['仿写', '小说', '故事']):
            return '仿写创作'
        elif any(word in filename for word in ['影视', '剧本', '电影']):
            return '影视小说'
        elif any(word in filename for word in ['营销', '策划']):
            return '营销策划'
        elif any(word in filename for word in ['职场', '面试', '规划']):
            return '职场办公'
        
        return '其他'

    def clean_title(self, filename):
        """清理文件名作为标题"""
        # 移除文件扩展名
        title = filename.replace('.docx', '').replace('.txt', '')
        
        # 移除常见的标记
        patterns = [
            r'【指令\+?.*?】',
            r'【指令】',
            r'【教程】',
            r'【视频教程】',
            r'\d+、?',
            r'^\d+\.',
        ]
        
        for pattern in patterns:
            title = re.sub(pattern, '', title)
        
        return title.strip()

    def extract_description_and_content(self, full_content, title):
        """从完整内容中提取描述和主要内容"""
        if not full_content:
            return "暂无描述", "暂无内容"
        
        lines = [line.strip() for line in full_content.split('\n') if line.strip()]
        
        if not lines:
            return "暂无描述", "暂无内容"
        
        # 如果内容很短，直接作为描述和内容
        if len(lines) <= 3:
            description = lines[0] if lines else "暂无描述"
            content = full_content
            return description, content
        
        # 寻找可能的描述部分
        description = ""
        content_start_idx = 0
        
        # 查找描述性文字
        for i, line in enumerate(lines[:5]):  # 只检查前5行
            if len(line) < 200 and not line.startswith(('Role:', 'Background:', 'Profile:')):
                if any(word in line for word in ['提示词', '指令', '助手', '专家', '生成', '创作']):
                    description = line
                    content_start_idx = i + 1
                    break
        
        # 如果没找到合适的描述，使用第一行
        if not description:
            description = lines[0]
            content_start_idx = 1
        
        # 提取主要内容
        content_lines = lines[content_start_idx:] if content_start_idx < len(lines) else lines
        content = '\n'.join(content_lines) if content_lines else full_content
        
        return description, content

    def parse_single_file(self, file_path):
        """解析单个文件"""
        logging.info(f"正在解析: {file_path}")
        
        # 提取内容
        if file_path.suffix == '.docx':
            full_content = self.extract_from_docx(file_path)
        else:
            full_content = self.extract_from_txt(file_path)
        
        if not full_content:
            self.stats['failed_files'].append(str(file_path))
            return None
        
        # 生成标题
        title = self.clean_title(file_path.name)
        
        # 确定分类
        category = self.determine_category(file_path)
        
        # 提取描述和内容
        description, content = self.extract_description_and_content(full_content, title)
        
        # 创建提示词对象
        prompt = {
            "提示词名称": title,
            "提示词描述": description,
            "提示词内容": content,
            "提示词分类": category
        }
        
        # 更新统计
        if category not in self.stats['categories']:
            self.stats['categories'][category] = 0
        self.stats['categories'][category] += 1
        
        self.stats['parsed_successfully'] += 1
        logging.info(f"解析成功: {title} -> {category}")
        
        return prompt

    def parse_all_files(self):
        """解析所有文件"""
        files = self.scan_all_files()
        
        logging.info("开始解析所有文件...")
        
        for file_path in files:
            prompt = self.parse_single_file(file_path)
            if prompt:
                self.prompts.append(prompt)
        
        logging.info(f"解析完成! 成功: {self.stats['parsed_successfully']}, 失败: {len(self.stats['failed_files'])}")

    def save_to_json(self):
        """保存到JSON文件"""
        data = {
            "prompts": self.prompts
        }
        
        with open(self.output_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        logging.info(f"数据已保存到: {self.output_file}")

    def print_statistics(self):
        """打印统计信息"""
        print("\n" + "="*50)
        print("📊 解析统计报告")
        print("="*50)
        print(f"总文件数: {self.stats['total_files']}")
        print(f"Word文档: {self.stats['docx_files']}")
        print(f"文本文件: {self.stats['txt_files']}")
        print(f"解析成功: {self.stats['parsed_successfully']}")
        print(f"解析失败: {len(self.stats['failed_files'])}")
        
        print(f"\n📂 分类统计:")
        for category, count in sorted(self.stats['categories'].items()):
            print(f"  {category}: {count}个")
        
        if self.stats['failed_files']:
            print(f"\n❌ 失败文件:")
            for file in self.stats['failed_files']:
                print(f"  {file}")

def main():
    source_dir = r"f:\个人文档\website\salary\2025豆包指令85+提示词合集"
    output_file = "prompts.json"
    
    parser = EnhancedPromptParser(source_dir, output_file)
    parser.parse_all_files()
    parser.save_to_json()
    parser.print_statistics()

if __name__ == "__main__":
    main()